"""Document loader: parse PDF, DOCX, TXT, MD, Image files with page tracking."""

from __future__ import annotations

import base64
import logging
from dataclasses import dataclass, field
from pathlib import Path

logger = logging.getLogger(__name__)


@dataclass
class PageText:
    page: int
    text: str


@dataclass
class Document:
    doc_id: str = field(default_factory=lambda: "")
    filename: str = ""
    text: str = ""
    file_type: str = ""
    file_size: int = 0
    metadata: dict = field(default_factory=dict)
    pages: list[PageText] = field(default_factory=list)
    generated_name: str = ""


IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp"}


class DocumentLoader:
    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".txt", ".md"} | IMAGE_EXTENSIONS

    def load(self, file_path: Path, doc_id: str = "") -> Document:
        ext = file_path.suffix.lower()
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {ext}")

        file_size = file_path.stat().st_size
        generated_name = ""

        if ext == ".pdf":
            pages = self._load_pdf(file_path)
            text = "\n\n".join(p.text for p in pages)
        elif ext in (".docx", ".doc"):
            pages = self._load_docx(file_path)
            text = "\n\n".join(p.text for p in pages)
        elif ext in IMAGE_EXTENSIONS:
            pages, generated_name = self._load_image(file_path)
            text = "\n\n".join(p.text for p in pages)
        else:
            pages = self._load_text(file_path)
            text = "\n\n".join(p.text for p in pages)

        return Document(
            doc_id=doc_id,
            filename=file_path.name,
            text=text,
            file_type=ext,
            file_size=file_size,
            pages=pages,
            generated_name=generated_name,
        )

    def _load_pdf(self, path: Path) -> list[PageText]:
        import pdfplumber

        pages: list[PageText] = []
        with pdfplumber.open(str(path)) as pdf:
            for i, page in enumerate(pdf.pages, start=1):
                page_text = page.extract_text()
                if page_text:
                    pages.append(PageText(page=i, text=page_text))
        return pages

    def _load_docx(self, path: Path) -> list[PageText]:
        from docx import Document as DocxDocument

        doc = DocxDocument(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n\n".join(paragraphs)
        per_page = max(1, len(paragraphs) // max(1, len(paragraphs) // 20))
        pages: list[PageText] = []
        page_num = 1
        for i in range(0, len(paragraphs), per_page):
            chunk = "\n\n".join(paragraphs[i : i + per_page])
            pages.append(PageText(page=page_num, text=chunk))
            page_num += 1
        if not pages and full_text.strip():
            pages.append(PageText(page=1, text=full_text))
        return pages

    def _load_image(self, path: Path) -> list[PageText]:
        from io import BytesIO

        from PIL import Image
        from openai import OpenAI

        from config import DASHSCOPE_API_KEY, KB_VISION_MODEL, KB_VISION_BASE_URL

        img = Image.open(str(path))
        if img.mode == "RGBA":
            img = img.convert("RGB")
        elif img.mode not in ("RGB", "L"):
            img = img.convert("RGB")

        buf = BytesIO()
        fmt = "PNG" if path.suffix.lower() in (".png", ".webp", ".bmp") else "JPEG"
        img.save(buf, format=fmt)
        b64 = base64.b64encode(buf.getvalue()).decode()
        mime = "image/png" if fmt == "PNG" else "image/jpeg"
        data_url = f"data:{mime};base64,{b64}"

        client = OpenAI(
            api_key=DASHSCOPE_API_KEY,
            base_url=KB_VISION_BASE_URL,
            timeout=120.0,
            max_retries=2,
        )
        completion = client.chat.completions.create(
            model=KB_VISION_MODEL,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "image_url", "image_url": {"url": data_url}},
                        {"type": "text", "text": "请提取图片中的所有文字内容。如果图片中没有文字，请简要描述图片的主要内容。"},
                    ],
                }
            ],
        )
        ocr_text = completion.choices[0].message.content or ""
        if not ocr_text.strip():
            return [], ""

        from config import LLM_API_KEY, LLM_BASE_URL, LLM_MODEL
        llm_client = OpenAI(
            api_key=LLM_API_KEY,
            base_url=LLM_BASE_URL,
            timeout=30.0,
            max_retries=2,
        )
        summary_completion = llm_client.chat.completions.create(
            model=LLM_MODEL,
            messages=[
                {"role": "system", "content": "你是一个标题生成器。根据给出的文本内容，生成一个简短的中文概要标题，不超过15个字，不要加标点，只输出标题本身。"},
                {"role": "user", "content": ocr_text.strip()},
            ],
        )
        summary_name = (summary_completion.choices[0].message.content or "").strip()[:15]

        return [PageText(page=1, text=ocr_text.strip())], summary_name

    def _load_text(self, path: Path) -> list[PageText]:
        encodings = ["utf-8", "gbk", "gb2312", "gb18030"]
        for enc in encodings:
            try:
                full_text = path.read_text(encoding=enc)
                lines = full_text.splitlines()
                per_page = 50
                pages: list[PageText] = []
                page_num = 1
                for i in range(0, len(lines), per_page):
                    pages.append(PageText(page=page_num, text="\n".join(lines[i : i + per_page])))
                    page_num += 1
                if not pages and full_text.strip():
                    pages.append(PageText(page=1, text=full_text))
                return pages
            except (UnicodeDecodeError, LookupError):
                continue
        raise RuntimeError(f"Cannot decode file: {path}")
